import requests

from docxtpl import DocxTemplate
from docx import Document
from openpyxl import load_workbook
from num2words import num2words
from pathlib import Path

wb = load_workbook('workdir/calc.xlsx')
page = wb["Page_1"]
data = {}


# Функция вывода словаря на экран
def print_data(data):
    for i, j in data.items():
        print(i, j)


# Получаем прописные значения и записываем их в словарь
def get_str_numbers(data):
    for j in data.values():
        j.append(num2words(j[0], lang='ru'))
        j.append(num2words(int(j[1]), lang='ru'))


# Создаем файлы и заполняем их данными
def make_files(data):
    for key, value in data.items():
        template_path = "workdir/example_2.docx"
        new_file_path = f"workdir/files/{key}.docx"
        year_coast = str(value[0])
        year_coast_str = value[3]
        days = str(value[2])
        period_coast_str = value[-1]
        period_coast_rub = str(int(value[1]))
        period_coast_penny = correct_penny(value[1])[-2:]
        correct_ruble_case_year = get_correct_ruble_case(year_coast)
        correct_ruble_case_period = get_correct_ruble_case(period_coast_rub)
        correct_penny_case_period = get_correct_penny(period_coast_penny)

        # extracted_text = extract_text_from_docx("workdir/example.docx")
        # extracted_text = extracted_text.replace("{№№}", key)
        # extracted_text = extracted_text.replace("{СУММ_Р}", period_coast_rub)
        # extracted_text = extracted_text.replace("{ПРОПИСЬ_Р}", period_coast_str)
        # extracted_text = extracted_text.replace("{РУБ_Р}", correct_ruble_case_period)
        # extracted_text = extracted_text.replace("{КОПЕЙКИ}", period_coast_penny)
        # extracted_text = extracted_text.replace("{КОП_Р_СКЛ}", correct_penny_case_period)
        # extracted_text = extracted_text.replace("{СУММ_Г}", year_coast)
        # extracted_text = extracted_text.replace("{ПРОПИСЬ_Г}", year_coast_str)
        # extracted_text = extracted_text.replace("{РУБ_Г}", correct_ruble_case_year)
        # extracted_text = extracted_text.replace("{КОЛВО_ДНЕЙ}", days)

        # new_doc = Document()
        # new_doc.add_paragraph(extracted_text)
        # new_doc.save(new_file_path)

        doc = DocxTemplate(template_path)
        context = {
            'key': key,
            'year_coast': year_coast,
            'year_coast_str': year_coast_str,
            'days': days,
            'period_coast_str': period_coast_str,
            'period_coast_rub': period_coast_rub,
            'period_coast_penny': period_coast_penny,
            'correct_ruble_case_year': correct_ruble_case_year,
            'correct_ruble_case_period': correct_ruble_case_period,
            'correct_penny_case_period': correct_penny_case_period,
        }
        doc.render(context)
        doc.save(new_file_path)

# Подбирает склонение рубля для числа
def get_correct_ruble_case(number: str) -> str:
    var1 = "056789"
    var2 = "234"
    if number[-1] in var1 or number[-2] == "1":
        return "рублей"
    elif number[-1] in var2:
        return "рубля"
    else:
        return "рубль"


# Подбирает склонение для копейки
def get_correct_penny(penny: str) -> str:
    var1 = "056789"
    var2 = "234"
    if penny[-1] in var1 or penny[-2] == "1":
        return "копеек"
    elif penny[-1] in var2:
        return "копейки"
    else:
        return "копейка"


# Функция для извлечения текста из файла Word
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def correct_penny(number: float) -> float:
    if str(number)[-2] == ".":
        return str(number) + "0"
    return str(number)


# Собираем данные из файла exal
for i in range(2, 66):
    key = str(page.cell(row=i, column=1).value)
    data[key] = [
        page.cell(row=i, column=4).value,
        round(page.cell(row=i, column=6).value, 2),
        page.cell(row=i, column=3).value,
    ]

get_str_numbers(data)
print_data(data)
make_files(data)

